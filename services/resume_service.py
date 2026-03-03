import hashlib
import io
import os
import re
import uuid
from pathlib import Path
from typing import List, Optional

from dotenv import load_dotenv
from fastapi import HTTPException
from openai import OpenAI
from sqlalchemy.orm import Session

from models.resume import Resume
from models.resume_keyword import ResumeKeyword
from models.resume_classification import ResumeClassification
from models.llm_run import LlmRun
from schemas.resume_llm import (
    ResumeClassificationResult,
    ResumeKeywordItem,
    ResumeKeywordResult,
)
from services.prompt.resume.classify_prompt import (
    PROMPT_VERSION_CLASSIFY,
    CLASSIFY_SYSTEM_PROMPT,
    build_classify_user_prompt,
)
from services.prompt.resume.keyword_prompt import (
    PROMPT_VERSION_KEYWORD,
    KEYWORD_SYSTEM_PROMPT,
    build_keyword_user_prompt,
)

load_dotenv()

DEFAULT_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")



PROJECT_ROOT = Path(__file__).resolve().parent.parent
UPLOAD_DIR = Path(
    os.environ.get("RESUME_UPLOAD_DIR", str(PROJECT_ROOT / "uploads" / "resume"))
)
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def normalize_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def safe_filename(filename: str) -> str:
    filename = os.path.basename(filename)
    return re.sub(r"[^a-zA-Z0-9._-가-힣]", "_", filename)


def detect_file_type(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return "PDF"
    if ext == ".docx":
        return "DOCX"

    raise HTTPException(
        status_code=400,
        detail="지원하지 않는 파일 형식입니다. PDF/DOCX만 허용합니다.",
    )


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()

def extract_pdf_text(data: bytes) -> str:
    try:
        import fitz
    except Exception as e:
        raise RuntimeError("pymupdf가 필요합니다. pip install pymupdf") from e

    doc = fitz.open(stream=data, filetype="pdf")
    return "\n".join(page.get_text("text") for page in doc)


def extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("python-docx가 필요합니다. pip install python-docx") from e

    file_obj = io.BytesIO(data)
    doc = Document(file_obj)

    parts: List[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_upload(filename: str, data: bytes) -> str:
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return extract_pdf_text(data)

    if ext == ".docx":
        return extract_docx_text(data)

    raise HTTPException(status_code=400, detail="지원하지 않는 파일 형식입니다.")


def get_client() -> OpenAI:
    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY가 없습니다. .env에 설정해 주세요.")
    return OpenAI(api_key=api_key)


def classify_resume_llm(
    resume_text: str,
    model: str = DEFAULT_MODEL,
) -> ResumeClassificationResult:
    client = get_client()

    resp = client.responses.parse(
        model=model,
        input=[
            {"role": "system", "content": CLASSIFY_SYSTEM_PROMPT},
            {"role": "user", "content": build_classify_user_prompt(resume_text)},
        ],
        text_format=ResumeClassificationResult,
        truncation="auto",
    )

    if resp.output_parsed is None:
        raise RuntimeError("이력서 분류 파싱 실패(output_parsed=None)")

    return resp.output_parsed


def analyze_resume_keywords_llm(
    resume_text: str,
    job_family: str,
    job_role: Optional[str],
    model: str = DEFAULT_MODEL,
) -> ResumeKeywordResult:
    client = get_client()

    resp = client.responses.parse(
        model=model,
        input=[
            {"role": "system", "content": KEYWORD_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": build_keyword_user_prompt(
                    resume_text=resume_text,
                    job_family=job_family,
                    job_role=job_role,
                ),
            },
        ],
        text_format=ResumeKeywordResult,
        truncation="auto",
    )

    if resp.output_parsed is None:
        raise RuntimeError("이력서 키워드 분석 파싱 실패(output_parsed=None)")

    return resp.output_parsed


def save_llm_run_success(
    db: Session,
    stage: str,
    model: str,
    prompt_version: str,
) -> LlmRun:
    row = LlmRun(
        llm_stage=stage,
        llm_model=model,
        llm_prompt_version=prompt_version,
        llm_status="SUCCESS",
        error_code=None,
        error_message=None,
    )
    db.add(row)
    db.flush()
    return row


def save_llm_run_failed(
    db: Session,
    stage: str,
    model: str,
    prompt_version: str,
    error_code: str,
    error_message: str,
) -> LlmRun:
    row = LlmRun(
        llm_stage=stage,
        llm_model=model,
        llm_prompt_version=prompt_version,
        llm_status="FAILED",
        error_code=error_code,
        error_message=error_message[:255] if error_message else None,
    )
    db.add(row)
    db.flush()
    return row


def dedupe_keywords(items: List[ResumeKeywordItem]) -> List[ResumeKeywordItem]:
    result: List[ResumeKeywordItem] = []
    seen = set()

    for item in items:
        keyword = (item.keyword or "").strip()
        if not keyword:
            continue

        key = (keyword.lower(), item.keyword_type)
        if key in seen:
            continue

        seen.add(key)
        result.append(item)

    return result


def create_resume_record(
    db: Session,
    user_id: int,
    original_filename: str,
    data: bytes,
) -> Resume:
    file_type = detect_file_type(original_filename)
    extracted_text = extract_text_from_upload(original_filename, data)
    extracted_text = normalize_text(extracted_text)

    if not extracted_text:
        raise HTTPException(
            status_code=400,
            detail="텍스트를 추출하지 못했습니다. 스캔 PDF일 수 있습니다.",
        )

    max_chars = 80000
    if len(extracted_text) > max_chars:
        extracted_text = extracted_text[:max_chars] + "\n\n[TRUNCATED]"

    resume = Resume(
        user_id=user_id,
        resume_file_name=original_filename,
        resume_file_type=file_type,
        resume_file_path=None,
        resume_file_size=len(data),
        resume_extracted_text=extracted_text,
        resume_sha256=sha256_bytes(data),
    )
    db.add(resume)
    db.commit()
    db.refresh(resume)
    return resume


def run_resume_classification(
    db: Session,
    resume: Resume,
    model: str = DEFAULT_MODEL,
):
    try:
        classification_result = classify_resume_llm(
            resume_text=resume.resume_extracted_text,
            model=model,
        )

        llm_run = save_llm_run_success(
            db=db,
            stage="RESUME_CLASSIFY",
            model=model,
            prompt_version=PROMPT_VERSION_CLASSIFY,
        )

        row = ResumeClassification(
            resume_id=resume.resume_id,
            llm_id=llm_run.llm_id,
            class_job_family=classification_result.job_family,
            class_job_role=classification_result.job_role,
            class_evidence=[item.model_dump() for item in classification_result.evidence],
        )
        db.add(row)
        db.commit()
        db.refresh(row)

        return row, classification_result

    except Exception as e:
        db.rollback()

        save_llm_run_failed(
            db=db,
            stage="RESUME_CLASSIFY",
            model=model,
            prompt_version=PROMPT_VERSION_CLASSIFY,
            error_code=type(e).__name__,
            error_message=str(e),
        )
        db.commit()

        raise HTTPException(
            status_code=500,
            detail=f"이력서 분류 실패: {type(e).__name__}: {e}",
        ) from e


def run_resume_keyword_analysis(
    db: Session,
    resume: Resume,
    classification: ResumeClassification,
    model: str = DEFAULT_MODEL,
):
    try:
        keyword_result = analyze_resume_keywords_llm(
            resume_text=resume.resume_extracted_text,
            job_family=classification.class_job_family,
            job_role=classification.class_job_role,
            model=model,
        )

        deduped = dedupe_keywords(keyword_result.keywords)

        llm_run = save_llm_run_success(
            db=db,
            stage="RESUME_KEYWORD",
            model=model,
            prompt_version=PROMPT_VERSION_KEYWORD,
        )

        for item in deduped:
            row = ResumeKeyword(
                resume_id=resume.resume_id,
                llm_id=llm_run.llm_id,
                keyword_keyword=item.keyword,
                keyword_type=item.keyword_type,
                keyword_evidence=[ev.model_dump() for ev in item.evidence],
            )
            db.add(row)

        db.commit()
        return deduped, keyword_result.notes

    except Exception as e:
        db.rollback()

        save_llm_run_failed(
            db=db,
            stage="RESUME_KEYWORD",
            model=model,
            prompt_version=PROMPT_VERSION_KEYWORD,
            error_code=type(e).__name__,
            error_message=str(e),
        )
        db.commit()

        raise HTTPException(
            status_code=500,
            detail=f"이력서 키워드 분석 실패: {type(e).__name__}: {e}",
        ) from e


def process_resume_upload_and_analyze(
    db: Session,
    user_id: int,
    original_filename: str,
    data: bytes,
    model: str = DEFAULT_MODEL,
):
    if not data:
        raise HTTPException(status_code=400, detail="빈 파일입니다.")

    resume = create_resume_record(
        db=db,
        user_id=user_id,
        original_filename=original_filename,
        data=data,
    )

    classification_row, classification_result = run_resume_classification(
        db=db,
        resume=resume,
        model=model,
    )

    keywords, keyword_notes = run_resume_keyword_analysis(
        db=db,
        resume=resume,
        classification=classification_row,
        model=model,
    )

    return {
        "resume_id": resume.resume_id,
        "file_name": resume.resume_file_name,
        "classification": {
            "job_family": classification_result.job_family,
            "job_role": classification_result.job_role,
            "evidence": [item.model_dump() for item in classification_result.evidence],
            "notes": classification_result.notes,
        },
        "keywords": [item.model_dump() for item in keywords],
        "keyword_notes": keyword_notes,
        "keyword_count": len(keywords),
    }